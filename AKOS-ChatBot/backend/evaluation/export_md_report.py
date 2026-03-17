from __future__ import annotations

import argparse
import re
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, Preformatted, SimpleDocTemplate, Spacer, Table, TableStyle


MARKDOWN_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def parse_markdown_lines(text: str):
    in_code = False
    for raw_line in text.splitlines():
        line = raw_line.rstrip("\n")

        if line.strip().startswith("```"):
            in_code = not in_code
            yield ("fence", "")
            continue

        if in_code:
            yield ("code", line)
            continue

        stripped = line.strip()
        if not stripped:
            yield ("blank", "")
        elif stripped.startswith("### "):
            yield ("h3", stripped[4:].strip())
        elif stripped.startswith("## "):
            yield ("h2", stripped[3:].strip())
        elif stripped.startswith("# "):
            yield ("h1", stripped[2:].strip())
        elif stripped.startswith("- ") or stripped.startswith("* "):
            yield ("bullet", stripped[2:].strip())
        elif stripped[:2].isdigit() and ". " in stripped[:4]:
            yield ("ordered", stripped)
        else:
            yield ("para", stripped)


def split_inline_segments(text: str) -> list[tuple[str, str, str | None]]:
    segments: list[tuple[str, str, str | None]] = []
    last_index = 0
    for match in MARKDOWN_LINK_RE.finditer(text):
        start, end = match.span()
        if start > last_index:
            segments.append(("text", text[last_index:start], None))
        label, href = match.group(1), match.group(2)
        segments.append(("link", label, href))
        last_index = end
    if last_index < len(text):
        segments.append(("text", text[last_index:], None))
    if not segments:
        segments.append(("text", text, None))
    return segments


def add_docx_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_properties.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_docx_markdown_paragraph(doc: Document, text: str, style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    for segment_type, segment_text, segment_target in split_inline_segments(text):
        if segment_type == "link" and segment_target:
            add_docx_hyperlink(paragraph, segment_text, segment_target)
        else:
            paragraph.add_run(segment_text)
    return paragraph


def markdown_to_pdf_markup(text: str) -> str:
    parts: list[str] = []
    for segment_type, segment_text, segment_target in split_inline_segments(text):
        if segment_type == "link" and segment_target:
            parts.append(
                f'<link href="{escape(segment_target)}" color="blue"><u>{escape(segment_text)}</u></link>'
            )
        else:
            parts.append(escape(segment_text))
    return "".join(parts)


def _split_table_row(line: str) -> list[str]:
    row = line.strip().strip("|")
    return [cell.strip() for cell in row.split("|")]


def _is_table_separator(line: str) -> bool:
    stripped = line.strip()
    if not (stripped.startswith("|") and stripped.endswith("|")):
        return False
    cells = _split_table_row(stripped)
    if not cells:
        return False
    for cell in cells:
        marker = cell.replace(":", "").replace("-", "")
        if marker != "":
            return False
    return True


def parse_markdown_blocks(text: str):
    lines = text.splitlines()
    i = 0
    in_code = False

    while i < len(lines):
        line = lines[i].rstrip("\n")
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            yield ("fence", "")
            i += 1
            continue

        if in_code:
            yield ("code", line)
            i += 1
            continue

        if (
            stripped.startswith("|")
            and i + 1 < len(lines)
            and _is_table_separator(lines[i + 1])
        ):
            headers = _split_table_row(stripped)
            i += 2
            rows: list[list[str]] = []
            while i < len(lines):
                candidate = lines[i].strip()
                if candidate.startswith("|") and candidate.endswith("|"):
                    rows.append(_split_table_row(candidate))
                    i += 1
                else:
                    break
            yield ("table", [headers, *rows])
            continue

        if not stripped:
            yield ("blank", "")
        elif stripped.startswith("### "):
            yield ("h3", stripped[4:].strip())
        elif stripped.startswith("## "):
            yield ("h2", stripped[3:].strip())
        elif stripped.startswith("# "):
            yield ("h1", stripped[2:].strip())
        elif stripped.startswith("- ") or stripped.startswith("* "):
            yield ("bullet", stripped[2:].strip())
        elif stripped[:2].isdigit() and ". " in stripped[:4]:
            yield ("ordered", stripped)
        else:
            yield ("para", stripped)

        i += 1


def export_docx(md_text: str, output_path: Path) -> None:
    doc = Document()

    for line_type, content in parse_markdown_blocks(md_text):
        if line_type == "blank" or line_type == "fence":
            continue
        if line_type == "h1":
            doc.add_heading(content, level=1)
        elif line_type == "h2":
            doc.add_heading(content, level=2)
        elif line_type == "h3":
            doc.add_heading(content, level=3)
        elif line_type == "bullet":
            add_docx_markdown_paragraph(doc, content, style="List Bullet")
        elif line_type == "ordered":
            add_docx_markdown_paragraph(doc, content, style="List Number")
        elif line_type == "code":
            doc.add_paragraph(content)
        elif line_type == "table":
            table_data: list[list[str]] = content
            max_cols = max(len(row) for row in table_data)
            table = doc.add_table(rows=len(table_data), cols=max_cols)
            table.style = "Table Grid"

            for row_idx, row_cells in enumerate(table_data):
                for col_idx in range(max_cols):
                    value = row_cells[col_idx] if col_idx < len(row_cells) else ""
                    cell = table.cell(row_idx, col_idx)
                    cell.text = ""
                    paragraph = cell.paragraphs[0]
                    for segment_type, segment_text, segment_target in split_inline_segments(value):
                        if segment_type == "link" and segment_target:
                            add_docx_hyperlink(paragraph, segment_text, segment_target)
                        else:
                            paragraph.add_run(segment_text)
                    if row_idx == 0:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("")
        else:
            add_docx_markdown_paragraph(doc, content)

    doc.save(output_path)


def _register_pdf_font() -> str:
    windows_arial = Path("C:/Windows/Fonts/arial.ttf")
    if windows_arial.exists():
        pdfmetrics.registerFont(TTFont("Arial", str(windows_arial)))
        return "Arial"
    return "Helvetica"


def export_pdf(md_text: str, output_path: Path) -> None:
    font_name = _register_pdf_font()

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "BodySl",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=10,
        leading=14,
    )
    h1 = ParagraphStyle(
        "H1Sl",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=18,
        leading=22,
        spaceAfter=10,
    )
    h2 = ParagraphStyle(
        "H2Sl",
        parent=styles["Heading2"],
        fontName=font_name,
        fontSize=14,
        leading=18,
        spaceAfter=8,
    )
    h3 = ParagraphStyle(
        "H3Sl",
        parent=styles["Heading3"],
        fontName=font_name,
        fontSize=12,
        leading=16,
        spaceAfter=6,
    )
    code = ParagraphStyle(
        "CodeSl",
        parent=body,
        fontName="Courier",
        fontSize=9,
        leading=12,
    )

    story = []
    in_code = False

    for line_type, content in parse_markdown_blocks(md_text):
        if line_type == "fence":
            in_code = not in_code
            story.append(Spacer(1, 4))
            continue

        if line_type == "blank":
            story.append(Spacer(1, 4))
            continue

        if in_code or line_type == "code":
            story.append(Preformatted(content, code))
            continue

        if line_type == "h1":
            safe = markdown_to_pdf_markup(content)
            story.append(Paragraph(safe, h1))
        elif line_type == "h2":
            safe = markdown_to_pdf_markup(content)
            story.append(Paragraph(safe, h2))
        elif line_type == "h3":
            safe = markdown_to_pdf_markup(content)
            story.append(Paragraph(safe, h3))
        elif line_type == "bullet":
            safe = markdown_to_pdf_markup(content)
            story.append(Paragraph(f"• {safe}", body))
        elif line_type == "table":
            table_data: list[list[str]] = content
            col_count = max(len(row) for row in table_data)
            normalized = [row + [""] * (col_count - len(row)) for row in table_data]
            paragraph_rows = []
            for row in normalized:
                paragraph_rows.append([Paragraph(markdown_to_pdf_markup(cell), body) for cell in row])
            pdf_table = Table(paragraph_rows, repeatRows=1)
            pdf_table.setStyle(
                TableStyle(
                    [
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EAEAEA")),
                        ("FONTNAME", (0, 0), (-1, 0), font_name),
                        ("FONTNAME", (0, 1), (-1, -1), font_name),
                        ("FONTSIZE", (0, 0), (-1, -1), 9),
                        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ]
                )
            )
            story.append(pdf_table)
            story.append(Spacer(1, 6))
        else:
            safe = markdown_to_pdf_markup(content)
            story.append(Paragraph(safe, body))

    pdf = SimpleDocTemplate(str(output_path), pagesize=A4)
    pdf.build(story)


def main() -> None:
    parser = argparse.ArgumentParser(description="Export Markdown report to DOCX and PDF.")
    parser.add_argument("--input", required=True, help="Path to Markdown file")
    parser.add_argument("--outdir", default=".", help="Output directory")
    args = parser.parse_args()

    input_path = Path(args.input).resolve()
    outdir = Path(args.outdir).resolve()
    outdir.mkdir(parents=True, exist_ok=True)

    md_text = input_path.read_text(encoding="utf-8")
    docx_path = outdir / f"{input_path.stem}.docx"
    pdf_path = outdir / f"{input_path.stem}.pdf"

    export_docx(md_text, docx_path)
    export_pdf(md_text, pdf_path)

    print(docx_path)
    print(pdf_path)


if __name__ == "__main__":
    main()
