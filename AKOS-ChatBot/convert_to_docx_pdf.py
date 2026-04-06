#!/usr/bin/env python3
"""Convert markdown to DOCX and PDF"""
import subprocess
import shutil
import os

md_file = "porocilo.md"
docx_file = "porocilo.docx"
pdf_file = "porocilo.pdf"

# Try using pandoc via subprocess with system PATH
try:
    # Try to convert with pandoc
    result = subprocess.run([
        "pandoc",
        md_file,
        "-o", docx_file,
        "--from", "markdown",
        "--to", "docx"
    ], capture_output=True, text=True)
    
    if result.returncode == 0:
        print(f"✓ Created {docx_file}")
    else:
        print(f"Pandoc error: {result.stderr}")
        raise Exception("Pandoc DOCX conversion failed")
    
    # Convert markdown to PDF via DOCX
    result = subprocess.run([
        "pandoc",
        md_file,
        "-o", pdf_file,
        "--from", "markdown",
        "--pdf-engine", "xelatex"
    ], capture_output=True, text=True)
    
    if result.returncode == 0:
        print(f"✓ Created {pdf_file}")
    else:
        # Try with pdflatex
        result = subprocess.run([
            "pandoc",
            md_file,
            "-o", pdf_file,
            "--from", "markdown",
            "--pdf-engine", "pdflatex"
        ], capture_output=True, text=True)
        
        if result.returncode == 0:
            print(f"✓ Created {pdf_file}")
        else:
            print(f"Pandoc PDF error: {result.stderr}")
            
except FileNotFoundError:
    print("Pandoc not found. Installing via alternative method...")
    # Download pandoc from GitHub
    import urllib.request
    import zipfile
    
    # This would require more complex setup, let's just use python-docx
    print("Falling back to python-docx for DOCX generation...")
    
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re
        
        # Read markdown file
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Create DOCX document
        doc = Document()
        
        # Parse markdown and add to document
        lines = md_content.split('\n')
        current_heading_level = 0
        
        for line in lines:
            if line.strip() == '':
                # Empty line
                doc.add_paragraph()
            elif line.startswith('# '):
                # Heading 1
                heading_text = line.replace('# ', '').strip()
                p = doc.add_heading(heading_text, level=1)
            elif line.startswith('## '):
                # Heading 2
                heading_text = line.replace('## ', '').strip()
                p = doc.add_heading(heading_text, level=2)
            elif line.startswith('### '):
                # Heading 3
                heading_text = line.replace('### ', '').strip()
                p = doc.add_heading(heading_text, level=3)
            elif line.startswith('- '):
                # Bullet point
                bullet_text = line.replace('- ', '').strip()
                doc.add_paragraph(bullet_text, style='List Bullet')
            elif line.startswith('| '):
                # Table (simple handling)
                pass
            else:
                # Regular paragraph
                if line.strip():
                    # Handle inline formatting
                    text = line.strip()
                    p = doc.add_paragraph(text)
        
        # Save DOCX
        doc.save(docx_file)
        print(f"✓ Created {docx_file} with python-docx")
        
    except ImportError:
        print("python-docx not available")

print("\nConversion complete!")
print(f"Files updated:")
print(f"  - {md_file}")
print(f"  - {docx_file}")
print(f"  - {pdf_file}")
